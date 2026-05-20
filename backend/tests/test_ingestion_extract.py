"""Extractor tests: PDF, DOCX, plain text, MIME dispatch."""

from __future__ import annotations

import pytest

from src.ingestion.extract import (
    detect_mime,
    extract,
    extract_docx,
    extract_pdf,
    extract_text,
)


def _make_test_pdf(path) -> None:
    import fitz

    doc = fitz.open()
    page = doc.new_page()
    page.insert_text((72, 72), "Adverse possession requires factual possession.")
    page = doc.new_page()
    page.insert_text((72, 72), "Page two with another paragraph of legal text.")
    doc.save(path)
    doc.close()


def _make_test_docx(path) -> None:
    import docx

    d = docx.Document()
    d.add_heading("Property Law Notes", level=1)
    d.add_paragraph("Adverse possession requires factual possession.")
    d.add_paragraph("Estoppel in equity prevents asserting strict rights.")
    d.save(str(path))


def test_extract_pdf_yields_page_sections(tmp_path):
    pdf = tmp_path / "test.pdf"
    _make_test_pdf(pdf)
    sections = extract_pdf(pdf)
    assert len(sections) == 2
    assert sections[0].page == 1
    assert "adverse possession" in sections[0].text.lower()
    assert sections[1].page == 2


def test_extract_pdf_skips_empty_pages(tmp_path):
    import fitz

    pdf = tmp_path / "with_blank.pdf"
    doc = fitz.open()
    doc.new_page()  # blank
    page = doc.new_page()
    page.insert_text((72, 72), "Some content here.")
    doc.save(pdf)
    doc.close()
    sections = extract_pdf(pdf)
    assert len(sections) == 1
    assert sections[0].page == 2


def test_extract_docx_returns_single_section(tmp_path):
    docx_path = tmp_path / "notes.docx"
    _make_test_docx(docx_path)
    sections = extract_docx(docx_path)
    assert len(sections) == 1
    assert "Property Law Notes" in sections[0].text
    assert "estoppel" in sections[0].text.lower()


def test_extract_text_handles_unicode(tmp_path):
    p = tmp_path / "notes.md"
    p.write_text("# Heading\n\nA note about *adverse* possession — émigré case.\n")
    sections = extract_text(p)
    assert len(sections) == 1
    assert "émigré" in sections[0].text


def test_extract_text_empty_returns_empty(tmp_path):
    p = tmp_path / "blank.txt"
    p.write_text("   \n  ")
    assert extract_text(p) == []


def test_detect_mime_extensions():
    assert detect_mime("file.pdf") == "application/pdf"
    assert detect_mime("file.docx") == (
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
    )
    assert detect_mime("note.md") == "text/markdown"
    assert detect_mime("a.png") == "image/png"
    assert detect_mime("unknown.foo987", fallback="application/octet-stream") == (
        "application/octet-stream"
    )


def test_extract_dispatches_by_mime(tmp_path):
    p = tmp_path / "notes.txt"
    p.write_text("hello")
    sections = extract(p)
    assert sections[0].text == "hello"


def test_extract_rejects_unknown_mime(tmp_path):
    p = tmp_path / "blob.bin"
    p.write_bytes(b"\x00\x01\x02")
    with pytest.raises(ValueError):
        extract(p)
